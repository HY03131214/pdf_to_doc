#!/usr/bin/env python3
"""
Generate Word document for chapters 1-3 (pages 9-61).
Builds document from scratch with template-matching styles.
"""

import json
import os
import re
import io
import hashlib
from PIL import Image
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import fitz

PDF_PATH = "/home/HY/pdf_to_doc/液压元件的寿命试验(段长宝等编)(2).pdf"
OCR_RESULT = "/home/HY/pdf_to_doc/ocr_results.json"
TEMPLATE_PATH = "/home/HY/pdf_to_doc/模板.docx"
OUTPUT_PATH = "/home/HY/pdf_to_doc/液压元件的寿命试验_第1-3章.docx"
IMG_DIR = "/home/HY/pdf_to_doc/temp_pages"

CHAPTER_RANGE = range(9, 62)  # pages 9-61 (chapters 1-3)

MATH_SYMBOLS = set('=+-×÷∫∑∏√∞∂αβγδελμστφωθηξρψΔΣΩΦΓ≥≤≠≈±·~→←↑↓⇒⇔∀∃∈∉⊂⊃∪∩∧∨¬∵∴∝∼≅≈≡⊆⊇⊕⊗⊥∥∠△□◇○◎◉◎°′″‰')

# Heading detection patterns
RE_CHAPTER = re.compile(r'^第[一二三四五六七八九十百\d]+章')
RE_SECTION = re.compile(r'^[\$§]?\d+\.\d+')
RE_HEADING_NUM = re.compile(r'^[一二三四五六七八九十]、')
RE_HEADING_PAREN = re.compile(r'^[（(][一二三四五六七八九十\d]+[）)]')
RE_HEADING_DOT = re.compile(r'^\d+\.\s+\S')
RE_FORMULA_CHAR = re.compile(r'[=≈≠<>±×÷∫∑∏√∂∞αβγδελμστφωθηξρ]')
RE_EQ_EQUATION = re.compile(r'^\s*[RrFfPpAaBbCcDdEeGgHhKkMmNnQqSsTtUuVvWwXxYyZz]\s*[=≈]')


def setup_styles(doc):
    """Configure document styles to match template."""
    # Page setup
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    # Normal style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.first_line_indent = Pt(10)
    style.paragraph_format.line_spacing = 1.3

    # Set East-Asian font for Normal
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = doc.styles['Normal'].element.get_or_add_rPr().makeelement(qn('w:rFonts'), {})
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')

    # Heading styles
    headings = [
        ('Heading 1', '黑体', Pt(18), WD_ALIGN_PARAGRAPH.CENTER, 0, None),
        ('Heading 2', '黑体', Pt(16), WD_ALIGN_PARAGRAPH.LEFT, 0, None),
        ('Heading 3', '黑体', Pt(15), WD_ALIGN_PARAGRAPH.LEFT, 0, None),
        ('Heading 4', '黑体', Pt(14), WD_ALIGN_PARAGRAPH.LEFT, 0, None),
    ]

    for name, font_name, font_size, align, indent, spacing in headings:
        if name in [s.name for s in doc.styles]:
            style = doc.styles[name]
        else:
            style = doc.styles.add_style(name, 1)  # PARAGRAPH type

        style.font.name = font_name
        style.font.size = font_size
        style.font.bold = False
        style.font.color.rgb = None
        style.paragraph_format.alignment = align
        style.paragraph_format.first_line_indent = indent if indent else 0
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(6)
        if spacing:
            style.paragraph_format.line_spacing = spacing

        # Set East-Asian font for heading
        rpr = style.element.get_or_add_rPr()
        rFonts_elem = rpr.find(qn('w:rFonts'))
        if rFonts_elem is None:
            rFonts_elem = rpr.makeelement(qn('w:rFonts'), {})
            rpr.insert(0, rFonts_elem)
        rFonts_elem.set(qn('w:eastAsia'), font_name)
        rFonts_elem.set(qn('w:ascii'), font_name)
        rFonts_elem.set(qn('w:hAnsi'), font_name)


def looks_like_number(text):
    """Check if text is just a number (possibly with spaces/parentheses)."""
    t = text.strip().replace(' ', '')
    # patterns like "0.0001", "1.204", "15.82-1.25s", "1.1~2", "70.20"
    if re.match(r'^[\d.~\-–—eE×x\s]+$', t):
        return True
    # patterns like "(8)", "(0)"
    if re.match(r'^[（(]\d+[）)]$', t):
        return True
    return False


def classify_line(text, y, x, page_lines, prev_y=None):
    """Classify a text line."""
    text = text.strip()
    if not text:
        return "skip"

    # Skip page numbers at top or bottom
    if re.match(r'^\d{1,3}$', text) and (y > 1400 or y < 80):
        return "skip"

    # Chapter heading
    if RE_CHAPTER.match(text):
        return "heading1"

    # Section heading like "1.1 Title", "2.1标题"
    # Must have small section numbers (e.g., "1.1" not "15.82")
    if RE_SECTION.match(text):
        sec_match = RE_SECTION.match(text)
        sec_num = sec_match.group()
        # Section numbers should be small (e.g., 1.1, 2.3, max 9.99)
        parts = sec_num.replace('$', '').replace('§', '').split('.')
        if len(parts) == 2:
            try:
                major = int(parts[0])
                if major > 10:  # unlikely to be a real section number
                    pass  # skip, treat as body/formula
                else:
                    after_num = text[text.find('.') + 1:].strip()
                    if len(after_num) >= 1 and not re.match(r'^[\d.~\-–—\s\-+]+$', after_num):
                        if len(text) > 5:
                            return "heading2"
            except ValueError:
                pass

    # Chinese numbered heading like "一、" — must be short and have meaningful content
    if RE_HEADING_NUM.match(text) and len(text) < 25:
        # Skip if it's just "一、" or "二、" without meaningful text
        if re.match(r'^[一二三四五六七八九十]、\s*$', text):
            pass  # too short, might be OCR error
        else:
            return "heading3"

    # Parenthesized heading like "（一）" or "(1)" — must be short and contain Chinese or be meaningful
    if RE_HEADING_PAREN.match(text) and len(text) < 25:
        # Exclude isolated equation numbers like "(8)", "(0)"
        if re.match(r'^[（(]\d+[）)]$', text):
            return "formula"
        # Exclude "(1)、(2)、(3)计算步骤同上" — this is body text referencing steps
        if len(text) > 15 and re.search(r'[，,、]', text):
            return "body"
        return "heading3"

    # Numbered item like "1. xxx" — must start near left margin
    if RE_HEADING_DOT.match(text) and x < 200 and len(text) < 40:
        return "heading4"

    # Catch numeric-looking false headings
    if looks_like_number(text):
        return "formula"

    # Formula detection
    math_count = sum(1 for c in text if c in MATH_SYMBOLS)
    if math_count >= 2 and len(text) < 80:
        return "formula"
    if RE_FORMULA_CHAR.search(text) and len(text) < 40:
        return "formula"
    if RE_EQ_EQUATION.search(text) and len(text) < 60:
        return "formula"

    return "body"


def determine_level(text, prev_level, x):
    """Determine heading level (1-4)."""
    text = text.strip()

    if RE_CHAPTER.match(text):
        return 1
    if RE_SECTION.match(text):
        return 2
    if RE_HEADING_NUM.match(text):
        return 2  # "一、" is typically Heading 2 in this book
    if RE_HEADING_PAREN.match(text):
        return 3
    if RE_HEADING_DOT.match(text):
        return 4
    if re.match(r'^\d+\.\s*$', text):  # just "1." - probably inline numbering
        return 4
    return 2


def insert_cropped_image(doc, page_img, bbox, scale=0.72):
    """Crop region from page and insert as centered image."""
    padding = 10
    x1 = max(0, int(bbox[0][0]) - padding)
    y1 = max(0, int(bbox[0][1]) - padding)
    x2 = min(page_img.width, int(bbox[2][0]) + padding)
    y2 = min(page_img.height, int(bbox[2][1]) + padding)

    if x2 - x1 < 5 or y2 - y1 < 5:
        return

    crop = page_img.crop((x1, y1, x2, y2))
    h = hashlib.md5(crop.tobytes()).hexdigest()[:10]
    crop_path = f"/home/HY/pdf_to_doc/temp_pages/_c_{h}.png"
    if not os.path.exists(crop_path):
        crop.save(crop_path)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run()
    width = Inches(crop.width * scale / 200)
    if width < Inches(0.3):
        width = Inches(0.3)
    run.add_picture(crop_path, width=width)


def insert_figure_region(doc, page_img, y1, y2, padding=20):
    """Insert a figure region (area between text lines) as image."""
    x1_pad = 100
    x2_pad = -100
    x_start = max(0, x1_pad)
    x_end = min(page_img.width, page_img.width + x2_pad)
    y_start = max(0, int(y1) - padding)
    y_end = min(page_img.height, int(y2) + padding)

    if y_end - y_start < 20 or x_end - x_start < 50:
        return

    crop = page_img.crop((x_start, y_start, x_end, y_end))
    h = hashlib.md5(crop.tobytes()).hexdigest()[:10]
    crop_path = f"/home/HY/pdf_to_doc/temp_pages/_g_{h}.png"
    if not os.path.exists(crop_path):
        crop.save(crop_path)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run()
    run.add_picture(crop_path, width=Inches(crop.width * 0.7 / 200))


def merge_body_lines(lines, page_img, doc):
    """Process a page's lines, merging body text and inserting formulas/figures."""
    i = 0
    n = len(lines)

    while i < n:
        ln = lines[i]
        text = ln["text"].strip()
        y = ln["y"]
        x = ln["x"]
        bbox = ln["bbox"]

        line_type = classify_line(text, y, x, lines)
        prev_y = lines[i - 1]["y"] if i > 0 else None

        if line_type == "skip":
            i += 1
            continue

        # Check for large vertical gap → possible figure
        if prev_y and line_type in ("body", "heading1", "heading2", "heading3", "heading4"):
            gap = y - prev_y
            if gap > 100:
                insert_figure_region(doc, page_img, prev_y, y)

        if line_type == "heading1":
            level = determine_level(text, 0, x)
            para = doc.add_paragraph(style=f"Heading {level}")
            para.add_run(text)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1

        elif line_type in ("heading2", "heading3", "heading4"):
            level = determine_level(text, 0, x)
            doc.add_paragraph(text, style=f"Heading {level}")
            i += 1

        elif line_type == "formula":
            # Group consecutive formulas
            fb_start = i
            while i < n and classify_line(lines[i]["text"].strip(), lines[i]["y"], lines[i]["x"], lines) == "formula":
                i += 1
            # Insert merged formula region
            f_bbox = [
                [lines[fb_start]["bbox"][0][0], lines[fb_start]["bbox"][0][1]],
                [lines[fb_start]["bbox"][1][0], lines[fb_start]["bbox"][1][1]],
                [lines[i - 1]["bbox"][2][0], lines[fb_start]["bbox"][0][1]],
                [lines[i - 1]["bbox"][3][0], lines[i - 1]["bbox"][2][1]],
            ]
            insert_cropped_image(doc, page_img, f_bbox)

        elif line_type == "body":
            # Collect consecutive body lines
            body_texts = []
            body_start = i
            while i < n:
                ln2 = lines[i]
                t2 = classify_line(ln2["text"].strip(), ln2["y"], ln2["x"], lines)

                if t2 not in ("body",):
                    break

                # Check for figure gap between body lines
                if body_texts:
                    gap = ln2["y"] - lines[i - 1]["y"]
                    if gap > 80:
                        # Insert the gap as figure and start new paragraph
                        insert_figure_region(doc, page_img, lines[i - 1]["y"], ln2["y"])
                        if body_texts:
                            para = doc.add_paragraph(style="Normal")
                            para.add_run("".join(body_texts))
                            body_texts = []
                        body_start = i

                body_texts.append(ln2["text"].strip())
                i += 1

            if body_texts:
                full = "".join(body_texts)
                para = doc.add_paragraph(style="Normal")
                para.add_run(full)

        else:
            i += 1


def main():
    print("Loading OCR results...")
    with open(OCR_RESULT, "r", encoding="utf-8") as f:
        ocr_data = json.load(f)

    print("Loading PDF...")
    pdf = fitz.open(PDF_PATH)

    print("Setting up document...")
    doc = Document()
    setup_styles(doc)

    for page_num in CHAPTER_RANGE:
        key = str(page_num)
        lines = ocr_data.get(key, [])
        if not lines:
            continue
        print(f"Page {page_num} ({len(lines)} lines)")

        pix = pdf[page_num - 1].get_pixmap(dpi=200)
        page_img = Image.open(io.BytesIO(pix.tobytes("png")))

        merge_body_lines(lines, page_img, doc)

    pdf.close()
    print(f"\nSaving to {OUTPUT_PATH}...")
    doc.save(OUTPUT_PATH)
    print("Done!")


if __name__ == "__main__":
    main()
